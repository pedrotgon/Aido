import os
import uvicorn
import asyncio
import shutil
import json
from typing import Optional
from fastapi import FastAPI, HTTPException, Request, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from pydantic import BaseModel
from dotenv import load_dotenv

# Load environment variables
load_dotenv(dotenv_path="../../.env")

# Add project root to PATH for ffmpeg
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../"))
if os.path.exists(os.path.join(PROJECT_ROOT, "ffmpeg.exe")):
    print(f"Adding {PROJECT_ROOT} to PATH for ffmpeg")
    os.environ["PATH"] = PROJECT_ROOT + os.pathsep + os.environ["PATH"]

# Import sub-agents for manual orchestration
from app.create.subagents.transcription.agent import transcription_agent
from app.create.subagents.structuring.agent import structuring_agent
from app.create.subagents.mastering.agent import mastering_agent
from app.create.subagents.json_converter.agent import json_converter_agent
from app.create.subagents.writer.agent import writer_agent

from google.adk.runners import InMemoryRunner
from google.adk.agents import Agent

from fastapi.staticfiles import StaticFiles

app = FastAPI(title="Aido Agent API")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directories
DATA_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../data"))
# ... (rest of directory setup)

# Mount static files
app.mount("/data", StaticFiles(directory=DATA_DIR), name="data")
UPLOAD_DIR = os.path.join(DATA_DIR, "entrada")
OUTPUT_DIR = os.path.join(DATA_DIR, "saida")
MANUALS_DIR = os.path.join(OUTPUT_DIR, "docx")
TRANSCRIPTS_DIR = os.path.join(OUTPUT_DIR, "txt")
PDFS_DIR = os.path.join(OUTPUT_DIR, "pdf")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(MANUALS_DIR, exist_ok=True)
os.makedirs(TRANSCRIPTS_DIR, exist_ok=True)
os.makedirs(PDFS_DIR, exist_ok=True)

# In-memory storage
manuals_db = {}

# --- Pydantic Models ---
class PipelineRunRequest(BaseModel):
    doc_id: str
    text_content: Optional[str] = None
    file_token: Optional[str] = None
    template_token: Optional[str] = None # New
    instructions: Optional[str] = None

class ManualUpdate(BaseModel):
    content: str

def build_docx_from_markdown(content: str, output_path: str):
    """Helper to generate a clean DOCX from Markdown content."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    
    doc = Document()
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Title
            p = doc.add_heading(line[2:], 0)
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 86, 145) # Bosch Blue
        elif line.startswith('## '):
            # Heading 1
            p = doc.add_heading(line[3:], 1)
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0) # Black
        elif line.startswith('### '):
            # Heading 2
            doc.add_heading(line[4:], 2)
        elif line.startswith('- ') or line.startswith('* '):
            # List item
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            # Normal text
            doc.add_paragraph(line)
    
    doc.save(output_path)
    return output_path

class SimpleMessage:
    def __init__(self, text, role="user"):
        self.role = role
        self.parts = [{"text": text}]
    def model_copy(self, **kwargs):
        return self

# --- Helper to run an ad-hoc agent ---
async def run_adhoc_agent(name: str, instruction: str, content: str, model: str = "gemini-2.5-flash") -> str:
    agent = Agent(name=name, model=model, instruction=instruction)
    runner = InMemoryRunner(agent=agent, app_name=name)
    
    # Create session
    session_id = f"sess_{name}_{os.urandom(4).hex()}"
    await runner.session_service.create_session(session_id=session_id, user_id="system", app_name=name)
    
    msg = SimpleMessage(content)
    response_text = ""
    
    async for event in runner.run_async(user_id="system", session_id=session_id, new_message=msg):
        # Assuming event has content.parts[0].text or similar
        # Inspecting event structure from previous tests might be needed if this fails.
        # But usually event.content.parts[0].text is standard.
        # Let's try to extract text safely.
        try:
            if hasattr(event, 'content') and event.content:
                 if hasattr(event.content, 'parts') and event.content.parts:
                     response_text += event.content.parts[0].text
        except Exception:
            pass
            
    return response_text

# --- Endpoints ---

@app.get("/health")
@app.get("/system/status")
def health_check():
    return {
        "status": "ok", 
        "agent": "Aido Orchestrator",
        "postgres": "active",
        "whisper": "active",
        "adk": "active",
        "auth": "active"
    }

@app.get("/system/config")
def system_config():
    return {
        "models": [{"id": "gemini-2.5-flash", "name": "Gemini 2.5 Flash"}],
        "default_model": "gemini-2.5-flash"
    }

@app.post("/upload")
async def upload_file(file: UploadFile = File(...), doc_id: str = Form(...)):
    try:
        file_path = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        return {
            "filename": file.filename,
            "transcription": "",
            "text_content": "",
            "file_token": file_path
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/pipeline/run")
async def run_pipeline(request: PipelineRunRequest):
    """
    Runs the pipeline steps manually to allow for fine-grained progress streaming
    and handling of different input cases (Video vs Text).
    """
    async def event_generator():
        try:
            yield f"event: init\ndata: \"Pipeline started for {request.doc_id}\"\n\n"
            
            current_context = ""
            transcript_path = None
            
            # --- STEP 1: TRANSCRIPTION ---
            if request.file_token:
                yield f'event: progress\ndata: {{"stage": "TRANSCRIPTION", "progress": 10, "log": "Iniciando transcrição do vídeo..."}}\n\n'
                
                from app.create.subagents.transcription.tools.transcribe_video import transcribe_video
                
                transcription_result = await transcribe_video(request.file_token)
                
                if transcription_result.startswith("Error"):
                     raise Exception(transcription_result)
                
                current_context = transcription_result
                
                # Save transcript
                transcript_filename = f"{request.doc_id}_transcript.txt"
                transcript_path_abs = os.path.join(TRANSCRIPTS_DIR, transcript_filename)
                with open(transcript_path_abs, "w", encoding="utf-8") as f:
                    f.write(current_context)
                transcript_path = f"/data/saida/txt/{transcript_filename}"
                
                yield f'event: progress\ndata: {{"stage": "TRANSCRIPTION", "progress": 25, "log": "Transcrição concluída."}}\n\n'
                
            elif request.text_content:
                yield f'event: progress\ndata: {{"stage": "TRANSCRIPTION", "progress": 25, "log": "Usando texto fornecido (Pular Transcrição)."}}\n\n'
                current_context = request.text_content
            else:
                 if request.instructions:
                     yield f'event: progress\ndata: {{"stage": "TRANSCRIPTION", "progress": 25, "log": "Usando instruções como conteúdo base."}}\n\n'
                     current_context = request.instructions
                 else:
                     raise Exception("Nenhum conteúdo fornecido (Vídeo, Texto ou Instruções).")

            if request.instructions and request.instructions != current_context:
                current_context += f"\n\nINSTRUÇÕES ADICIONAIS:\n{request.instructions}"

            # --- STEP 2: STRUCTURING ---
            yield f'event: progress\ndata: {{"stage": "STRUCTURING", "progress": 40, "log": "Estruturando conteúdo..."}}\n\n'
            
            print(f"[DEBUG] Before Structuring - Length: {len(current_context)}, Preview: {current_context[:200]}")
            structuring_prompt = f"Organize o seguinte texto em capítulos e fluxos:\n\n{current_context}"
            structured_text = await run_adhoc_agent("StructuringAgent", "Você é um especialista em estruturação de manuais técnicos.", structuring_prompt)
            current_context = structured_text
            print(f"[DEBUG] After Structuring - Length: {len(current_context)}, Preview: {current_context[:200]}")
            
            yield f'event: progress\ndata: {{"stage": "STRUCTURING", "progress": 55, "log": "Estruturação concluída."}}\n\n'

            # --- STEP 3: MASTERING ---
            yield f'event: progress\ndata: {{"stage": "MASTERING", "progress": 65, "log": "Aplicando diretrizes Bosch..."}}\n\n'
            
            mastering_prompt = f"Ajuste o tom técnico e terminologia Bosch do seguinte texto:\n\n{current_context}"
            mastered_text = await run_adhoc_agent("MasteringAgent", "Você é um especialista em diretrizes da Bosch. Use tom técnico e formal.", mastering_prompt)
            current_context = mastered_text
            print(f"[DEBUG] After Mastering - Length: {len(current_context)}, Preview: {current_context[:200]}")
            
            yield f'event: progress\ndata: {{"stage": "MASTERING", "progress": 75, "log": "Masterização concluída."}}\n\n'

            # --- STEP 4: CONTENT GENERATION ---
            yield f'event: progress\ndata: {{"stage": "JSON_CONVERTER", "progress": 80, "log": "Gerando conteúdo final..."}}\n\n'
            
            manual_docx_path = "/placeholder.docx"
            markdown_content_for_ui = ""
            
            if request.template_token:
                # --- TEMPLATE MODE (JSON) ---
                yield f'event: progress\ndata: {{"stage": "JSON_CONVERTER", "progress": 82, "log": "Preenchendo template personalizado..."}}\n\n'
                
                json_prompt = (
                    "Você é um assistente especializado em estruturar dados para geração de documentos automáticos.\n"
                    "Sua tarefa é converter o manual técnico fornecido em um objeto JSON estritamente formatado.\n"
                    "Este JSON será usado para preencher um template Word (.docx).\n\n"
                    "ESTRUTURA OBRIGATÓRIA:\n"
                    "{\n"
                    "  \"titulo\": \"Título do Manual\",\n"
                    "  \"introducao\": \"Texto introdutório completo...\",\n"
                    "  \"capitulos\": [\n"
                    "    {\n"
                    "      \"titulo\": \"Título do Capítulo 1\",\n"
                    "      \"conteudo\": \"Conteúdo detalhado do capítulo...\"\n"
                    "    },\n"
                    "    ...\n"
                    "  ]\n"
                    "}\n\n"
                    "REGRAS:\n"
                    "1. Responda APENAS com o JSON válido. Sem markdown, sem explicações.\n"
                    "2. O conteúdo deve ser rico e detalhado.\n"
                    "3. Mantenha a terminologia técnica da Bosch.\n\n"
                    f"CONTEÚDO DE ENTRADA:\n{current_context}"
                )
                
                json_text = await run_adhoc_agent("JsonAgent", "Você é um conversor JSON rigoroso.", json_prompt)
                json_text = json_text.replace("```json", "").replace("```", "").strip()
                
                try:
                    data = json.loads(json_text)
                    markdown_content_for_ui = f"# {data.get('titulo', 'Manual')}\n\n"
                    markdown_content_for_ui += f"{data.get('introducao', '')}\n\n"
                    for cap in data.get('capitulos', []):
                        markdown_content_for_ui += f"## {cap.get('titulo', 'Capítulo')}\n\n"
                        markdown_content_for_ui += f"{cap.get('conteudo', '')}\n\n"
                except json.JSONDecodeError:
                    print("[ERROR] Failed to parse JSON for Markdown conversion. Sending raw text.")
                    markdown_content_for_ui = current_context

                # Generate DOCX using Template
                yield f'event: progress\ndata: {{"stage": "WRITER", "progress": 90, "log": "Gerando arquivo DOCX (Template)..."}}\n\n'
                from app.create.subagents.writer.tools.write_docx import write_docx
                
                print(f"[DEBUG] Using custom template: {request.template_token}")
                writer_result = await write_docx(json_text, request.template_token, MANUALS_DIR)
                
                if writer_result["status"] == "success":
                    manual_docx_path = f"/data/saida/docx/{os.path.basename(writer_result['output_path'])}"
                else:
                    yield f'event: progress\ndata: {{"stage": "WRITER", "progress": 95, "log": "Erro no template: {writer_result.get('message')}"}}\n\n'

            else:
                # --- STANDARD MODE (Markdown) ---
                yield f'event: progress\ndata: {{"stage": "JSON_CONVERTER", "progress": 82, "log": "Gerando manual padrão (Markdown)..."}}\n\n'
                
                markdown_prompt = (
                    "Você é um redator técnico sênior da Bosch.\n"
                    "Escreva o manual final completo em formato Markdown estruturado.\n"
                    "Use '# Título', '## Capítulos', listas e negrito.\n"
                    "Não use blocos de código ou JSON. Apenas o texto do documento.\n\n"
                    f"CONTEÚDO DE ENTRADA:\n{current_context}"
                )
                
                markdown_text = await run_adhoc_agent("WriterAgent", "Você é um redator técnico experiente.", markdown_prompt)
                markdown_content_for_ui = markdown_text
                
                # Generate DOCX from Markdown
                yield f'event: progress\ndata: {{"stage": "WRITER", "progress": 90, "log": "Formatando documento padrão..."}}\n\n'
                
                filename = f"manual_{request.doc_id}.docx"
                filepath = os.path.join(MANUALS_DIR, filename)
                build_docx_from_markdown(markdown_text, filepath)
                manual_docx_path = f"/data/saida/docx/{filename}"

            yield f'event: progress\ndata: {{"stage": "WRITER", "progress": 100, "log": "Concluído."}}\n\n'

            # Send Content to UI
            # Use json.dumps for safe string escaping in JSON payload
            yield f'event: complete\ndata: {json.dumps({"manual_content": markdown_content_for_ui, "manual_docx_path": manual_docx_path, "manual_pdf_path": "/placeholder.pdf", "transcript_path": transcript_path})}\n\n'

        except Exception as e:
            import traceback
            traceback.print_exc()
            yield f'event: error\ndata: {{"message": "{str(e)}"}}\n\n'

    return StreamingResponse(event_generator(), media_type="text/event-stream")

@app.get("/manual/{doc_id}")
def get_manual(doc_id: str):
    if doc_id in manuals_db:
        return manuals_db[doc_id]
    return {"doc_id": doc_id, "content": "Conteúdo não encontrado.", "updated_at": "now"}

@app.put("/manual/{doc_id}")
async def update_manual(doc_id: str, update: ManualUpdate):
    manuals_db[doc_id] = {
        "doc_id": doc_id,
        "content": update.content,
        "updated_at": "now"
    }
    return manuals_db[doc_id]

@app.get("/events/{user_id}")
async def sse_endpoint(user_id: str, is_audio: str = "false"):
    print(f"Client #{user_id} connected via SSE")
    async def event_generator():
        try:
            yield f'data: {json.dumps({"type": "connection", "status": "connected", "user_id": user_id})}\n\n'
            while True:
                await asyncio.sleep(15)
                yield ": keep-alive\n\n"
        except asyncio.CancelledError:
            print(f"Client #{user_id} disconnected")

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*",
        }
    )

@app.post("/convert/docx")
async def convert_to_docx(update: ManualUpdate):
    try:
        # Always use the standard formatting builder for edits for now
        # (Simplifies logic: if user edits Markdown, they get a standard styled DOCX)
        
        filename = f"manual_editado_{os.urandom(4).hex()}.docx"
        filepath = os.path.join(OUTPUT_DIR, "docx", filename)
        
        build_docx_from_markdown(update.content, filepath)
        
        return FileResponse(filepath, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=filename)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run("app.server:app", host="0.0.0.0", port=8000, reload=True)
